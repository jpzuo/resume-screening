#!/usr/bin/env python3
"""Extract PDF and DOCX resumes into an auditable batch manifest."""

from __future__ import annotations

import argparse
import hashlib
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def normalize_text(value: str | None) -> str:
    return " ".join((value or "").split())


def extract_pdf(path: Path) -> tuple[list[dict[str, str]], str]:
    reader = PdfReader(str(path))
    sources: list[dict[str, str]] = []
    for index, page in enumerate(reader.pages, start=1):
        text = normalize_text(page.extract_text())
        if text:
            sources.append({"location": f"第{index}页", "text": text})
    if not sources:
        raise ValueError("扫描 PDF 或无可提取文本")
    return sources, "\n".join(item["text"] for item in sources)


def extract_docx(path: Path) -> tuple[list[dict[str, str]], str]:
    document = Document(str(path))
    sources: list[dict[str, str]] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = normalize_text(paragraph.text)
        if text:
            sources.append({"location": f"第{index}段", "text": text})
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            text = normalize_text(" | ".join(cell.text for cell in row.cells))
            if text:
                sources.append({"location": f"表格{table_index}第{row_index}行", "text": text})
    if not sources:
        raise ValueError("DOCX 未包含可提取文本")
    return sources, "\n".join(item["text"] for item in sources)


def candidate_id(relative_path: Path, index: int) -> str:
    digest = hashlib.sha1(str(relative_path).encode("utf-8")).hexdigest()[:6].upper()
    return f"R{index:03d}-{digest}"


def discover_resumes(input_dir: Path) -> list[Path]:
    files = [
        path
        for path in input_dir.rglob("*")
        if path.is_file() and not path.is_symlink() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    return sorted(files, key=lambda path: str(path.relative_to(input_dir)).lower())


def extract_document(path: Path) -> tuple[list[dict[str, str]], str]:
    if path.suffix.lower() == ".pdf":
        return extract_pdf(path)
    return extract_docx(path)


def build_manifest(input_dir: Path) -> dict[str, Any]:
    documents: list[dict[str, Any]] = []
    for index, path in enumerate(discover_resumes(input_dir), start=1):
        relative_path = path.relative_to(input_dir)
        document: dict[str, Any] = {
            "candidate_id": candidate_id(relative_path, index),
            "relative_path": relative_path.as_posix(),
            "file_name": path.name,
            "file_extension": path.suffix.lower().lstrip(".").upper(),
            "status": "已解析",
            "parse_error": "",
            "sources": [],
            "extracted_text": "",
        }
        try:
            sources, extracted_text = extract_document(path)
            document["sources"] = sources
            document["extracted_text"] = extracted_text
        except Exception as error:  # Report per-file failures without losing the batch.
            document["status"] = "解析异常"
            document["parse_error"] = str(error)
        documents.append(document)
    if not documents:
        raise ValueError("简历文件夹中未找到 PDF 或 DOCX 文件")
    return {
        "schema_version": "1.0",
        "input_dir": str(input_dir),
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "documents": documents,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="提取 PDF/DOCX 简历到批次清单")
    parser.add_argument("--input-dir", required=True, help="简历文件夹")
    parser.add_argument("--manifest", required=True, help="写入批次清单 JSON 的位置")
    args = parser.parse_args()

    input_dir = Path(args.input_dir).expanduser().resolve()
    manifest_path = Path(args.manifest).expanduser().resolve()
    if not input_dir.is_dir():
        raise SystemExit(f"简历文件夹不存在或不是目录: {input_dir}")
    if manifest_path.is_relative_to(input_dir):
        raise SystemExit("批次清单不能写入简历文件夹")

    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(
        json.dumps(build_manifest(input_dir), ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(manifest_path)


if __name__ == "__main__":
    main()
